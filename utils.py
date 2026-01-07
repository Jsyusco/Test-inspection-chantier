# utils.py
import streamlit as st
import pandas as pd
import numpy as np
import uuid
import json
import zipfile
import io
import urllib.parse
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL

# --- CONSTANTES ---
PROJECT_RENAME_MAP = {
    'Intitulé': 'Intitulé',
    'Fournisseur Bornes AC [Bornes]': 'Fournisseur Bornes AC',
    'Fournisseur Bornes DC [Bornes]': 'Fournisseur Bornes DC',
    'L [Plan de Déploiement]': 'PDC Lent',
    'R [Plan de Déploiement]': 'PDC Rapide',
    'UR [Plan de Déploiement]': 'PDC Ultra-rapide',
    'Pré L [Plan de Déploiement]': 'PDC L pré-équipés',
    'Pré R [Plan de Déploiement]': 'PDC R pré-équipés',
    'Pré UR [Plan de Déploiement]': 'PDC UR pré-équipés',
}

DISPLAY_GROUPS = [
    ['Intitulé', 'Fournisseur Bornes AC [Bornes]', 'Fournisseur Bornes DC [Bornes]'],
    ['L [Plan de Déploiement]', 'R [Plan de Déploiement]', 'UR [Plan de Déploiement]'],
    ['Pré L [Plan de Déploiement]', 'Pré R [Plan de Déploiement]', 'Pré UR [Plan de Déploiement]']
]

# --- FONCTIONS DE CHARGEMENT ---

def load_data_from_gsheets():
    """
    Charge les données via l'URL publique exportée en CSV pour éviter l'erreur 400.
    """
    try:
        # Récupération de l'URL depuis les secrets
        base_url = st.secrets["connections"]["gsheets"]["spreadsheet"]
        # Nettoyage de l'URL pour extraire l'ID
        if "/edit" in base_url:
            sheet_id = base_url.split("/d/")[1].split("/edit")[0]
        else:
            sheet_id = "1D-LInL839f997L7V_S394749_O3-4-8" # ID de secours

        # URLs d'export CSV pour chaque onglet
        url_questions = f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet=Questions"
        url_sites = f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet=Sites"

        df_questions = pd.read_csv(url_questions)
        df_sites = pd.read_csv(url_sites)

        # Nettoyage basique pour éviter les colonnes fantômes (NaN)
        df_questions = df_questions.dropna(how='all', axis=1).dropna(how='all', axis=0)
        df_sites = df_sites.dropna(how='all', axis=1).dropna(how='all', axis=0)

        return df_questions, df_sites
    except Exception as e:
        st.error(f"Erreur de lecture Google Sheets : {e}")
        return None, None

# --- GÉNÉRATION DU RAPPORT WORD ---

def generate_word_report(df_struct, project_data, start_time):
    doc = Document()
    
    # Style de titre
    title = doc.add_heading('Rapport d\'Inspection de Chantier', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Infos projet
    doc.add_heading('Informations Générales', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    info_dict = {
        "Projet": project_data.get('Intitulé', 'N/A'),
        "Date": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "Délai d'inspection": f"{datetime.now() - start_time}".split('.')[0]
    }
    
    for k, v in info_dict.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(k)
        row_cells[1].text = str(v)
        
    # Réponses (Logique simplifiée pour l'exemple)
    doc.add_heading('Détails de l\'Audit', level=1)
    for index, row in df_struct.iterrows():
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{row['Question']}: ").bold = True
        # On récupère la réponse si elle existe dans le session_state (géré par app.py)
        # Note : cette partie dépend de comment vous stockez vos réponses finales
        
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- INTERFACE DE QUESTION ---

def render_question(q_id, q_text, q_type, q_options, q_desc, q_mandatory, answers, phase_name, key_suffix="", loop_index=0, is_dynamic_comment=False):
    """
    Affiche un composant de question Streamlit selon son type.
    """
    label_html = f"<strong>{q_id}. {q_text}</strong>" + (' <span class="mandatory">*</span>' if q_mandatory else "")
    widget_key = f"q_{q_id}_{phase_name}_{key_suffix}_{loop_index}"
    current_val = answers.get(q_id)

    st.markdown(f'<div class="question-card"><div>{label_html}</div>', unsafe_allow_html=True)
    if q_desc and not pd.isna(q_desc): 
        st.markdown(f'<div class="description">⚠️ {q_desc}</div>', unsafe_allow_html=True)
    
    if q_type == 'text':
        if is_dynamic_comment:
            answers[q_id] = st.text_area("Réponse", value=current_val if current_val else "", key=widget_key, label_visibility="collapsed")
        else:
            answers[q_id] = st.text_input("Réponse", value=current_val if current_val else "", key=widget_key, label_visibility="collapsed")
            
    elif q_type == 'select':
        if isinstance(q_options, str):
            opts = [o.strip() for o in q_options.split(',')]
        else:
            opts = []
        if "" not in opts: opts.insert(0, "")
        answers[q_id] = st.selectbox("Sélectionnez", opts, index=opts.index(current_val) if current_val in opts else 0, key=widget_key, label_visibility="collapsed")
        
    elif q_type == 'number':
        try:
            val = int(current_val) if current_val else 0
        except:
            val = 0
        answers[q_id] = st.number_input("Nombre", value=val, step=1, key=widget_key, label_visibility="collapsed")
        
    elif q_type == 'photo':
        answers[q_id] = st.file_uploader("📸 Joindre photo(s)", accept_multiple_files=True, key=widget_key)
        
    st.markdown('</div>', unsafe_allow_html=True)
